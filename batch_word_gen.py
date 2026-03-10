import os
import re
import threading
import tkinter as tk
from tkinter import filedialog, messagebox, ttk

import pandas as pd
from docx import Document


def build_mapping(row: pd.Series) -> dict:
    """
    同时支持两种占位符：
    1) __字段A__  
    2) {字段A}
    """
    mapping = {}
    for col, val in row.items():
        v = "" if pd.isna(val) else str(val)
        mapping[f"__{col}__"] = v
        mapping[f"{{{col}}}"] = v
    return mapping


# def replace_in_paragraph(paragraph, mapping: dict):
#     """
#     为了尽量保留样式：逐 run 替换（比整段重写更稳）
#     注意：如果占位符被 Word 自动拆到多个 run，run 级别替换可能漏。
#     """
#     for run in paragraph.runs:
#         text = run.text
#         new_text = text
#         for k, v in mapping.items():
#             if k in new_text:
#                 new_text = new_text.replace(k, v)
#         if new_text != text:
#             run.text = new_text


def replace_in_paragraph(paragraph, mapping: dict):
    """
    按整段文本替换
    """
    full_text = "".join(run.text for run in paragraph.runs)
    new_text = full_text

    for k, v in mapping.items():
        new_text = new_text.replace(k, v)

    if new_text != full_text:
        # 清空原 runs
        for run in paragraph.runs:
            run.text = ""
        # 写回到第一个 run；如果没有 run，就新建一个
        if paragraph.runs:
            paragraph.runs[0].text = new_text
        else:
            paragraph.add_run(new_text)


def replace_in_doc(doc: Document, mapping: dict):
    # 普通段落
    for p in doc.paragraphs:
        replace_in_paragraph(p, mapping)

    # 表格内段落
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for p in cell.paragraphs:
                    replace_in_paragraph(p, mapping)

# def replace_in_doc(doc: Document, mapping: dict):
#     # 段落
#     for p in doc.paragraphs:
#         replace_in_paragraph(p, mapping)

#     # 表格
#     for table in doc.tables:
#         for row in table.rows:
#             for cell in row.cells:
#                 for p in cell.paragraphs:
#                     replace_in_paragraph(p, mapping)


class App:
    def __init__(self, root: tk.Tk):
        self.root = root
        root.title("批量Word生成工具")
        root.geometry("780x440")

        self.template_path = tk.StringVar()
        self.excel_path = tk.StringVar()
        self.out_dir = tk.StringVar()
        self.status = tk.StringVar(value="准备就绪")
        self.filename_col = tk.StringVar(value="字段A")  # 默认用字段A当文件名

        frm = ttk.Frame(root, padding=12)
        frm.pack(fill=tk.BOTH, expand=True)

        # Word 模板
        ttk.Label(frm, text="Word模板(.docx)：").grid(row=0, column=0, sticky=tk.W)
        ttk.Entry(frm, textvariable=self.template_path, width=62).grid(row=0, column=1, padx=6)
        ttk.Button(frm, text="选择模板", command=self.pick_template).grid(row=0, column=2)

        # Excel
        ttk.Label(frm, text="Excel数据(.xlsx)：").grid(row=1, column=0, sticky=tk.W, pady=(6, 0))
        ttk.Entry(frm, textvariable=self.excel_path, width=62).grid(row=1, column=1, padx=6, pady=(6, 0))
        ttk.Button(frm, text="选择Excel", command=self.pick_excel).grid(row=1, column=2, pady=(6, 0))

        # 输出目录
        ttk.Label(frm, text="输出目录：").grid(row=2, column=0, sticky=tk.W, pady=(6, 0))
        ttk.Entry(frm, textvariable=self.out_dir, width=62).grid(row=2, column=1, padx=6, pady=(6, 0))
        ttk.Button(frm, text="选择目录", command=self.pick_outdir).grid(row=2, column=2, pady=(6, 0))

        # 文件名列
        ttk.Label(frm, text="文件名用哪一列：").grid(row=3, column=0, sticky=tk.W, pady=(10, 0))
        ttk.Entry(frm, textvariable=self.filename_col, width=20).grid(row=3, column=1, sticky=tk.W, padx=6, pady=(10, 0))

        # 字段预览
        ttk.Label(frm, text="Excel字段预览：").grid(row=4, column=0, sticky=tk.W, pady=(12, 0))
        self.listbox = tk.Listbox(frm, height=8, width=66)
        self.listbox.grid(row=5, column=0, columnspan=2, sticky=tk.W, pady=6)

        # 按钮
        btns = ttk.Frame(frm)
        btns.grid(row=6, column=0, columnspan=3, pady=10, sticky=tk.W)
        ttk.Button(btns, text="刷新字段", command=self.load_fields).pack(side=tk.LEFT, padx=6)
        ttk.Button(btns, text="开始生成", command=self.start).pack(side=tk.LEFT, padx=6)

        ttk.Label(frm, textvariable=self.status).grid(row=7, column=0, columnspan=3, sticky=tk.W, pady=(10, 0))

    def pick_template(self):
        p = filedialog.askopenfilename(filetypes=[("Word", "*.docx")])
        if p:
            self.template_path.set(p)
            self.status.set(f"已选择模板：{os.path.basename(p)}")

    def pick_excel(self):
        p = filedialog.askopenfilename(filetypes=[("Excel", "*.xlsx;*.xls")])
        if p:
            self.excel_path.set(p)
            self.status.set(f"已选择Excel：{os.path.basename(p)}")
            self.load_fields()

    def pick_outdir(self):
        d = filedialog.askdirectory()
        if d:
            self.out_dir.set(d)

    def load_fields(self):
        self.listbox.delete(0, tk.END)
        ep = self.excel_path.get()
        if not ep or not os.path.exists(ep):
            self.status.set("请先选择Excel")
            return
        try:
            df = pd.read_excel(ep, dtype=str)
            for c in df.columns.tolist():
                self.listbox.insert(tk.END, c)
            self.status.set(f"Excel字段读取成功：{len(df.columns)}列，{len(df)}行数据")
        except Exception as e:
            messagebox.showerror("错误", f"读取Excel失败：{e}")

    def start(self):
        t = self.template_path.get()
        e = self.excel_path.get()
        out = self.out_dir.get()

        if not t or not os.path.exists(t):
            messagebox.showwarning("提示", "请选择Word模板(.docx)")
            return
        if not e or not os.path.exists(e):
            messagebox.showwarning("提示", "请选择Excel数据")
            return
        if not out:
            messagebox.showwarning("提示", "请选择输出目录")
            return

        threading.Thread(target=self.worker, daemon=True).start()

    def worker(self):
        try:
            self.status.set("开始生成...")
            df = pd.read_excel(self.excel_path.get(), dtype=str).fillna("")
            total = len(df)
            if total == 0:
                messagebox.showinfo("提示", "Excel没有数据行")
                self.status.set("无数据")
                return

            name_col = self.filename_col.get().strip()
            if name_col and name_col not in df.columns:
                self.status.set(f"文件名列“{name_col}”不存在，改用默认 output_序号")
                name_col = ""

            for idx, row in df.iterrows():
                mapping = build_mapping(row)

                doc = Document(self.template_path.get())
                replace_in_doc(doc, mapping)

                # 文件名
                if name_col:
                    base = str(row[name_col]).strip() or f"output_{idx+1}"
                else:
                    base = f"output_{idx+1}"

                # Windows非法字符处理
                base = re.sub(r'[\\/:*?"<>|]+', "_", base).strip()
                out_path = os.path.join(self.out_dir.get(), f"{base}.docx")

                # 避免覆盖
                if os.path.exists(out_path):
                    out_path = os.path.join(self.out_dir.get(), f"{base}_{idx+1}.docx")

                doc.save(out_path)
                self.status.set(f"已生成 {idx+1}/{total}：{os.path.basename(out_path)}")

            messagebox.showinfo("完成", f"全部生成完成，共 {total} 个文件。")
            self.status.set("生成完成")
        except Exception as e:
            messagebox.showerror("错误", str(e))
            self.status.set("出错：" + str(e))


if __name__ == "__main__":
    root = tk.Tk()
    App(root)
    root.mainloop()
